#-*-coding: utf-8-*-

# 依檔號選擇科室

import pandas as pd
excel_name = '106D000233_236to239_263.xlsx' # excel檔名
df = pd.read_excel(excel_name)
newdf = df[df.檔號.str.match('CLIO06')] # 科室編號
len(newdf.index)


from docx import Document  # 使用 docx 模組
document = Document() #創建空的 word template
section = document.sections[0]


from docx.shared import Cm, Pt  #加入可調整的 word 單位


#調整文件左右上下邊界至 1.27 cm
section.left_margin=Cm(3.17)
section.right_margin=Cm(3.17)
section.top_margin=Cm(1.54)
section.bottom_margin=Cm(1.54)


#把A4文件轉成橫置方向

from docx.enum.section import WD_ORIENT #處理文件的直向/橫向
section.orientation = WD_ORIENT.LANDSCAPE
new_width, new_height = Cm(29.7), Cm(21)
section.page_width = new_width
section.page_height= new_height


#加入字串並調整中文為標楷體
from docx.oxml.ns import qn


def create_page_by_row(row_number):
    
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'檔案銷毀目錄（案卷）')
    font = run.font
    font.name= 'New Times Roman'   #設定英文字體
    font.size=Pt(16)
    font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  #設定中文字體

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'勞動部職業安全衛生署 檔 案 銷 毀 目 錄')
    font = run.font
    font.name= 'New Times Roman'
    font.size=Pt(16)
    font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  #設定中文字體

    批號 = newdf.iat[row_number,1] # 批號列
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'批號：	' + 批號)
    font = run.font
    font.name= 'New Times Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  #設定中文字體

    年度 = str(newdf.iat[row_number,2]) # 年度列
    檔號 = str(newdf.iat[row_number,3]) # 檔號列
    案次 = str(newdf.iat[row_number,4]) # 案次列
    卷數 = str(newdf.iat[row_number,5]) # 卷數列
    起迄日期 = str(newdf.iat[row_number,6]) # 案卷日期起迄列
    保存年限 = str(newdf.iat[row_number,7]) # 保存年限列
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'檔號：' + 年度 + '/' + 檔號 + '/' + 案次 + '	卷數：' + 卷數 + '		案卷內文件起迄日期：' + 起迄日期 + '		保存年限：' + 保存年限)
    font = run.font
    font.name= 'New Times Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    案名 = newdf.iat[row_number,8] # 案名列
    檔案產生者 = newdf.iat[row_number,9] # 檔案產生者列
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'案名：	' + 案名 +'				檔案產生者：	' + 檔案產生者 + '	調整後保存年限(調整原因)：')
    font = run.font
    font.name= 'New Times Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  

    案情摘要 = newdf.iat[row_number,10] # 案情摘要列
    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'案情摘要：' + 案情摘要)
    font = run.font
    font.name= 'New Times Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'備註：原檔案產生機關為原行政院勞工委員會中區勞動檢查所')
    font = run.font
    font.name= 'New Times Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') 

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'銷毀檔案總數量：')
    font = run.font
    font.name= 'New Times Roman'
    font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  
    run = paragraph.add_run(u'	1	案	' + 卷數 + '	卷：')
    font = run.font
    font.name= 'New Times Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'			承辦人：				簽章			監毀人：				簽章')
    font = run.font
    font.name= 'New Times Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體') 

    paragraph = document.add_paragraph()
    run = paragraph.add_run(u'銷　毀　作　業')
    font = run.font
    font.name= 'New Times Roman'
    font.bold = True
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  
    run = paragraph.add_run(u'	核准銷毀文號：					銷毀日期：')
    font = run.font
    font.name= 'New Times Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')  


for i in range(len(newdf.index)):
    create_page_by_row(i)
    document.add_page_break()

#檔案存檔
document.save('秘書室.docx')
